from docx import Document

def generate_docx(student, role, template="professional"):
    filename = f"static/{student['name']}_{role}.docx"
    doc = Document()

    doc.add_heading(student["name"], level=1)
    doc.add_paragraph(f"{student['email']} | {student['phone']}")

    if student.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(student["summary"])

    doc.add_heading("Skills", level=2)
    for s in student["skills"]:
        doc.add_paragraph(s.strip(), style="List Bullet")

    doc.add_heading("Education", level=2)
    for edu in student["education"]:
        doc.add_paragraph(edu)

    doc.add_heading("Projects", level=2)
    for proj in student["projects"]:
        doc.add_paragraph(proj, style="List Bullet")

    doc.add_heading("Certifications", level=2)
    for c in student["certifications"]:
        doc.add_paragraph(c, style="List Bullet")

    if student.get("declaration"):
        doc.add_heading("Declaration", level=2)
        doc.add_paragraph(student["declaration"])

    doc.save(filename)
    return filename
